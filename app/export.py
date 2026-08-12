"""Export functionality for PDF, HTML, and DOCX generation."""

import os
import logging
import re as _re
from io import BytesIO
from typing import Tuple
try:
    from weasyprint import HTML
except (OSError, ImportError):
    HTML = None

GTK3_MISSING_MSG = (
    "PDF export requires GTK3. Download the installer: "
    "https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/"
    "releases/download/2022-01-04/gtk3-runtime-3.24.31-2022-01-04-ts-win64.exe"
)
from flask import Blueprint, render_template, request, send_file, jsonify, Response
from markupsafe import Markup
from .models import Proposal
from .utils import build_export_context, build_tracker_export_context
from .config import Config, ERROR_MESSAGES

logger = logging.getLogger(__name__)

export_bp = Blueprint("export", __name__)

EXPORT_DIR = Config.EXPORTS_DIR


def ensure_export_dir() -> None:
    """Ensure export directory exists."""
    os.makedirs(EXPORT_DIR, exist_ok=True)


def _md_to_plain(text: str) -> str:
    """Strip markdown to plain text for DOCX."""
    if not text:
        return ""
    text = _re.sub(r'#{1,6}\s+', '', text)
    text = _re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = _re.sub(r'\*(.+?)\*', r'\1', text)
    text = _re.sub(r'`(.+?)`', r'\1', text)
    text = _re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = _re.sub(r'^[-*]\s+', '  - ', text, flags=_re.MULTILINE)
    return text.strip()


def _add_page_number_footer(doc) -> None:
    """Add a right-aligned, small gray page number to the document footer."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_timeline_to_docx(doc, proposal) -> None:
    """Render the Gantt-style timeline for DOCX exports.

    Mirrors the Timeline section shown in the PDF/HTML exports: a landscape
    page with a year header row, a month header row, and shaded bars per
    task/budget item (dark blue for tasks, light blue for budget items).
    """
    from datetime import datetime as _dt
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION, WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _shade_cell(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    ctx = build_export_context(proposal)
    timeline_rows = ctx.get("all_rows") or []
    if not timeline_rows:
        return

    total_months = ctx.get("timeline_total_months", 1)
    if total_months < 1:
        total_months = 1

    try:
        sd = _dt.strptime(proposal.start_date, "%Y-%m-%d")
        proj_start_month = sd.month
        proj_start_year = sd.year
    except (ValueError, TypeError):
        proj_start_month = 1
        proj_start_year = 2025

    month_names = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
                   'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']

    # Landscape section for the chart (the PDF export uses a landscape page too).
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = Cm(1.2)
    section.left_margin = section.right_margin = Cm(1.5)

    doc.add_heading('Timeline', level=1)

    ncols = total_months + 1
    table = doc.add_table(rows=2, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.5)
    usable_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm - 4.5
    month_w = Cm(max(usable_cm / max(total_months, 1), 1.0))
    for i in range(total_months):
        table.columns[i + 1].width = month_w

    # Year header row (merged across each year's months, like the HTML export).
    years_in_proj = (proj_start_month - 1 + total_months - 1) // 12 + 1
    for y in range(years_in_proj):
        abs_y = proj_start_year + y
        year_start_m = max((abs_y - proj_start_year) * 12 - (proj_start_month - 1), 0)
        year_end_m = min((y + 1) * 12 - (proj_start_month - 1), total_months)
        year_span = year_end_m - year_start_m
        if year_span <= 0:
            continue
        first = table.rows[0].cells[year_start_m + 1]
        last = table.rows[0].cells[year_start_m + year_span]
        cell = first if year_span == 1 else first.merge(last)
        cell.text = str(abs_y)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Month header row.
    corner = table.rows[1].cells[0]
    corner.text = 'Task'
    for run in corner.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
    for m in range(total_months):
        cell = table.rows[1].cells[m + 1]
        cell.text = month_names[(proj_start_month - 1 + m) % 12]
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8)

    # Data rows: one shaded bar per activity, matching the HTML/PDF bar colors.
    for r in timeline_rows:
        row = table.add_row()
        name_cell = row.cells[0]
        name_cell.text = r["name"]
        if r.get("is_indent"):
            name_cell.paragraphs[0].paragraph_format.left_indent = Cm(0.4)
        for run in name_cell.paragraphs[0].runs:
            run.font.size = Pt(9)

        fill = '#bfdbfe' if r.get("is_indent") else '#1d4ed8'
        for bar in r["bars"]:
            start = max(int(bar.get("offset", 0)), 0)
            dur = int(bar.get("duration", 1))
            end = min(start + dur, total_months)
            if end <= start:
                continue
            first = row.cells[start + 1]
            last = row.cells[end]
            cell = first if end == start + 1 else first.merge(last)
            _shade_cell(cell, fill)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _build_proposal_docx(proposal) -> BytesIO:
    """Build a DOCX document for a proposal."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_page_number_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(proposal.title or "Proposal")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    if proposal.client_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Funder: {proposal.client_name}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if proposal.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Program: {proposal.subtitle}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(proposal.updated_at[:10] if proposal.updated_at else "")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    if proposal.project_summary:
        doc.add_heading('Project Summary', level=1)
        for line in _md_to_plain(proposal.project_summary).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    if proposal.scope:
        doc.add_heading('Scope', level=1)
        for line in _md_to_plain(proposal.scope).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    map_config = getattr(proposal, 'map_config', None) or {}
    if map_config.get("show_in_preview"):
        from .utils import build_geolibre_embed_src
        doc.add_heading('Map', level=1)
        image_url = (map_config.get("image_url") or "").strip()
        if image_url:
            try:
                import urllib.request
                with urllib.request.urlopen(image_url, timeout=15) as resp:
                    img_buf = BytesIO(resp.read())
                doc.add_picture(img_buf, width=Inches(6))
            except Exception:
                logger.warning(f"Could not fetch map image for DOCX export: {image_url}")
        p = doc.add_paragraph()
        p.add_run("Interactive map: ").bold = True
        p.add_run(build_geolibre_embed_src(map_config))
        p = doc.add_paragraph()
        p.add_run("Map powered by GeoLibre — Wu, Q. (2026). GeoLibre: A lightweight, "
                  "cloud-native GIS platform for visualizing, exploring, and analyzing "
                  "geospatial data. Zenodo. https://doi.org/10.5281/zenodo.20785400")
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    if proposal.budget_items:
        doc.add_heading('Budget', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['Item', 'Cost/Unit', 'Units', 'Total']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for task in proposal.tasks:
            task_items = [b for b in proposal.budget_items if b.get("task_id") == task["id"]]
            if not task_items:
                continue
            task_total = sum(i.get("cost_per_unit", 0) * i.get("units", 0) for i in task_items)
            row = table.add_row()
            row.cells[0].text = task["name"]
            row.cells[3].text = f"{task_total:,.0f}"
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            for item in task_items:
                total = item.get("cost_per_unit", 0) * item.get("units", 0)
                row = table.add_row()
                row.cells[0].text = f"  {item['name']}"
                row.cells[1].text = f"{item.get('cost_per_unit', 0):,.0f}"
                row.cells[2].text = f"{int(item.get('units', 0))}"
                row.cells[3].text = f"{total:,.0f}"
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

        row = table.add_row()
        row.cells[0].text = 'Subtotal'
        row.cells[3].text = f"{proposal.total_budget:,.0f}"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        indirect_percent = getattr(proposal, 'indirect_percent', 0) or 0
        if indirect_percent > 0:
            indirect_amount = proposal.total_budget * (indirect_percent / 100)
            row = table.add_row()
            row.cells[0].text = f"Indirect ({indirect_percent:.0f}%)"
            row.cells[3].text = f"{indirect_amount:,.0f}"

        row = table.add_row()
        row.cells[0].text = 'Total'
        indirect_amount = proposal.total_budget * (indirect_percent / 100)
        row.cells[3].text = f"${proposal.total_budget + indirect_amount:,.0f}"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        from .utils import build_budget_by_year
        by_year = build_budget_by_year(proposal)
        if by_year["years"] or by_year["unscheduled"]:
            year_table = doc.add_table(rows=1, cols=3)
            year_table.style = 'Light Grid Accent 1'
            year_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(['Year', 'Amount', '% of Total']):
                cell = year_table.rows[0].cells[i]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            total_budget = proposal.total_budget
            for row in by_year["years"]:
                r = year_table.add_row()
                r.cells[0].text = str(row["year"])
                r.cells[1].text = f"{row['amount']:,.0f}"
                pct = f"{(row['amount'] / total_budget * 100):.1f}%" if total_budget > 0 else ""
                r.cells[2].text = pct
            for u in by_year["unscheduled"]:
                r = year_table.add_row()
                r.cells[0].text = 'Not scheduled'
                r.cells[1].text = f"{u['amount']:,.0f}"
                pct = f"{(u['amount'] / total_budget * 100):.1f}%" if total_budget > 0 else ""
                r.cells[2].text = pct
            for r in year_table.rows[1:]:
                for cell in r.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    if proposal.show_budget_description and proposal.budget_description:
        doc.add_heading('Budget Description', level=1)
        for line in _md_to_plain(proposal.budget_description).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    if proposal.qualifications:
        doc.add_heading('Qualifications', level=1)
        for line in _md_to_plain(proposal.qualifications).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    if proposal.custom_sections:
        for section in sorted(proposal.custom_sections, key=lambda s: s.get("order", 0)):
            doc.add_heading(section.get("title", "Section"), level=1)
            for line in _md_to_plain(section.get("content", "")).split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

    _add_timeline_to_docx(doc, proposal)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _build_tracker_docx(proposal, ctx) -> BytesIO:
    """Build a DOCX document for the project tracker."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_page_number_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Project Report: {proposal.title}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    if proposal.client_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{proposal.client_name}{(' - ' + proposal.subtitle) if proposal.subtitle else ''}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    doc.add_heading('Dashboard Summary', level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Tasks Completed', 'Total Tasks', 'Total Budget', 'Milestones']
    values = [f"{ctx['overall_pct']}%", str(len(ctx['tasks'])), f"${ctx['total_with_indirect']:,.0f}", str(len(ctx['milestones']))]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for i, v in enumerate(values):
        table.rows[1].cells[i].text = v
        for run in table.rows[1].cells[i].paragraphs[0].runs:
            run.font.size = Pt(10)

    status_labels = {
        "not_started": "Not Started",
        "in_progress": "In Progress",
        "completed": "Completed",
        "delayed": "Delayed",
    }

    doc.add_heading('Tasks', level=1)
    for task in ctx['tasks']:
        doc.add_heading(task.get('name', 'Untitled'), level=2)
        status = status_labels.get(task.get('status', 'not_started'), 'Not Started')
        progress = task.get('progress_pct', 0)
        p = doc.add_paragraph()
        run = p.add_run(f"Status: {status}")
        run.bold = True
        p = doc.add_paragraph(f"Progress: {progress}%")
        if task.get('actual_start'):
            p = doc.add_paragraph(f"Actual Start: {task['actual_start']}")
        if task.get('actual_end'):
            p = doc.add_paragraph(f"Actual End: {task['actual_end']}")
        if task.get('notes'):
            doc.add_paragraph()
            run = p.add_run("Notes:")
            run.bold = True
            for line in task['notes'].split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

    doc.add_heading('Budget Tracking', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Task', 'Item', 'Planned', 'Actual', 'Variance']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for task in ctx['tasks']:
        tid = task.get("id", "")
        if tid not in ctx['task_budgets']:
            continue
        tb = ctx['task_budgets'][tid]
        row = table.add_row()
        row.cells[0].text = task.get('name', '')
        row.cells[2].text = f"${tb['subtotal']:,.0f}"
        row.cells[3].text = f"${tb['actual_total']:,.0f}"
        variance = tb['actual_total'] - tb['subtotal']
        row.cells[4].text = f"${variance:,.0f}"
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for item in tb["items"]:
            planned = item.get("cost_per_unit", 0) * item.get("units", 0)
            actual = item.get("actual_cost", 0)
            v = actual - planned
            row = table.add_row()
            row.cells[1].text = item.get('name', '')
            row.cells[2].text = f"${planned:,.0f}"
            row.cells[3].text = f"${actual:,.0f}"
            row.cells[4].text = f"${v:,.0f}"
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

    row = table.add_row()
    row.cells[0].text = 'Total'
    row.cells[2].text = f"${ctx['total_with_indirect']:,.0f}"
    row.cells[3].text = f"${ctx['total_actual'] + ctx['indirect_amount']:,.0f}"
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    if ctx['milestones']:
        doc.add_heading('Milestones', level=1)
        for m in ctx['milestones']:
            check = "[x]" if m.get("completed") else "[ ]"
            date_str = f" ({m['date']})" if m.get('date') else ''
            doc.add_paragraph(f"{check} {m.get('name', '')}{date_str}")

    if ctx['reports']:
        doc.add_heading('Reports', level=1)
        for r in reversed(ctx['reports']):
            doc.add_heading(r.get('title', 'Report'), level=2)
            if r.get('date'):
                p = doc.add_paragraph(f"Date: {r['date']}")
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            if r.get('content'):
                for line in r['content'].split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@export_bp.route("/export/pdf/<proposal_id>")
def export_pdf(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export proposal as PDF."""
    if HTML is None:
        return jsonify({"error": GTK3_MISSING_MSG}), 500

    proposal = Proposal.load(proposal_id)
    if not proposal:
        logger.warning(f"Proposal not found for PDF export: {proposal_id}")
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_export_context(proposal)
    html_content = render_template("export_proposal.html", **ctx)

    ensure_export_dir()
    pdf_path = os.path.join(EXPORT_DIR, f"{proposal_id}.pdf")
    HTML(string=html_content, base_url=request.host_url).write_pdf(pdf_path)

    return send_file(
        pdf_path,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{proposal.title or 'proposal'}.pdf",
    )


@export_bp.route("/export/html/<proposal_id>")
def export_html(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export proposal as HTML file."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        logger.warning(f"Proposal not found for HTML export: {proposal_id}")
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_export_context(proposal)
    html_content = render_template("export_proposal.html", **ctx)

    return Response(
        html_content,
        mimetype="text/html",
        headers={"Content-Disposition": "inline"},
    )


@export_bp.route("/export/docx/<proposal_id>")
def export_docx(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export proposal as DOCX file."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        logger.warning(f"Proposal not found for DOCX export: {proposal_id}")
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    try:
        buf = _build_proposal_docx(proposal)
    except ImportError:
        return jsonify({"error": "python-docx not installed"}), 500
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        return jsonify({"error": f"DOCX export failed: {str(e)}"}), 500

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{proposal.title or 'proposal'}.docx",
    )


@export_bp.route("/export/timeline/png/<proposal_id>")
def export_timeline_png(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export just the timeline chart as a PNG image.

    Renders the chart with WeasyPrint (same engine as the PDF export) onto a
    content-sized page, then rasterizes it with pypdfium2 so the image is a
    tight crop rather than a full page.
    """
    if HTML is None:
        return jsonify({"error": GTK3_MISSING_MSG}), 500

    try:
        import pypdfium2 as pdfium
    except ImportError:
        return jsonify({"error": "pypdfium2 not installed"}), 500

    proposal = Proposal.load(proposal_id)
    if not proposal:
        logger.warning(f"Proposal not found for timeline PNG export: {proposal_id}")
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_export_context(proposal)
    html_content = render_template("timeline_export.html", **ctx)

    ensure_export_dir()
    pdf_buf = BytesIO()
    HTML(string=html_content, base_url=request.host_url).write_pdf(pdf_buf)
    pdf_buf.seek(0)

    pdf = pdfium.PdfDocument(pdf_buf)
    page = pdf[0]
    bitmap = page.render(scale=2)
    pil = bitmap.to_pil()
    png_path = os.path.join(EXPORT_DIR, f"timeline_{proposal_id}.png")
    pil.save(png_path, "PNG")

    return send_file(
        png_path,
        mimetype="image/png",
        as_attachment=True,
        download_name=f"{proposal.title or 'proposal'}_timeline.png",
    )


@export_bp.route("/export/tracker/pdf/<proposal_id>")
def export_tracker_pdf(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export tracker as PDF."""
    if HTML is None:
        return jsonify({"error": GTK3_MISSING_MSG}), 500

    proposal = Proposal.load(proposal_id)
    if not proposal:
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_tracker_export_context(proposal)
    html_content = render_template("export_tracker.html", **ctx)

    ensure_export_dir()
    pdf_path = os.path.join(EXPORT_DIR, f"tracker_{proposal_id}.pdf")
    HTML(string=html_content, base_url=request.host_url).write_pdf(pdf_path)

    return send_file(
        pdf_path,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{proposal.title or 'project'}_tracker.pdf",
    )


@export_bp.route("/export/tracker/html/<proposal_id>")
def export_tracker_html(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export tracker as HTML."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_tracker_export_context(proposal)
    html_content = render_template("export_tracker.html", **ctx)

    return Response(
        html_content,
        mimetype="text/html",
        headers={"Content-Disposition": "inline"},
    )


@export_bp.route("/export/tracker/docx/<proposal_id>")
def export_tracker_docx(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export tracker as DOCX."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_tracker_export_context(proposal)

    try:
        buf = _build_tracker_docx(proposal, ctx)
    except ImportError:
        return jsonify({"error": "python-docx not installed"}), 500
    except Exception as e:
        logger.error(f"Tracker DOCX export failed: {e}")
        return jsonify({"error": f"DOCX export failed: {str(e)}"}), 500

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{proposal.title or 'project'}_tracker.docx",
    )
